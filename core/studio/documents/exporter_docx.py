"""DOCX exporter — render DocumentContentTree to a Word document."""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from core.schemas.studio_schema import DocumentContentTree, DocumentSection


# XML-invalid control characters (except tab, newline, carriage return)
_CTRL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _sanitize(text: str) -> str:
    """Strip XML-invalid control characters."""
    return _CTRL_CHAR_RE.sub("", text)


def export_to_docx(
    content_tree: DocumentContentTree,
    output_path: Path,
) -> Path:
    """Render a DocumentContentTree to a DOCX file.

    Uses python-docx with default Word styles.
    Returns the output path.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Set default font size
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading(_sanitize(content_tree.doc_title), level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Abstract
    if content_tree.abstract:
        doc.add_paragraph("")  # spacer
        abstract_para = doc.add_paragraph(_sanitize(content_tree.abstract))
        abstract_para.style = doc.styles["Normal"]
        # Italicize abstract
        for run in abstract_para.runs:
            run.italic = True
        doc.add_paragraph("")  # spacer after abstract

    # Sections
    _render_sections(doc, content_tree.sections)

    # Bibliography
    if content_tree.bibliography:
        doc.add_heading("Bibliography", level=1)
        for entry in content_tree.bibliography:
            key = entry.get("key", "")
            title = entry.get("title", "")
            author = entry.get("author", "")
            bib_text = f"[{key}] {author}. {title}."
            # Add extra fields if present
            year = entry.get("year", "")
            url = entry.get("url", "")
            if year:
                bib_text += f" ({year})"
            if url:
                bib_text += f" {url}"
            doc.add_paragraph(_sanitize(bib_text))

    doc.save(str(output_path))
    return output_path


def _render_sections(doc: Document, sections: List[DocumentSection]) -> None:
    """Recursively render sections as headings + paragraphs."""
    for section in sections:
        # Heading level: 1-3 (Word supports up to heading 9, but we cap at 3)
        heading_level = min(section.level, 3)
        doc.add_heading(_sanitize(section.heading), level=heading_level)

        # Content paragraphs
        if section.content:
            # Split on double newlines for paragraph breaks
            paragraphs = section.content.split("\n\n")
            for para_text in paragraphs:
                text = _sanitize(para_text.strip())
                if text:
                    # Render inline citations as [key]
                    doc.add_paragraph(text)

        # Recurse into subsections
        if section.subsections:
            _render_sections(doc, section.subsections)
