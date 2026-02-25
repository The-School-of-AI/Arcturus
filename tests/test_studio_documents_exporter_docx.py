"""Tests for core/studio/documents/exporter_docx.py — DOCX rendering."""

import pytest
from pathlib import Path

from docx import Document

from core.schemas.studio_schema import DocumentContentTree, DocumentSection
from core.studio.documents.exporter_docx import export_to_docx


@pytest.fixture
def sample_content_tree():
    return DocumentContentTree(
        doc_title="Test Report",
        doc_type="report",
        abstract="This is a test abstract for the report.",
        sections=[
            DocumentSection(
                id="sec1",
                heading="Introduction",
                level=1,
                content="Introduction paragraph one.\n\nIntroduction paragraph two.",
                subsections=[
                    DocumentSection(
                        id="sec1a",
                        heading="Background",
                        level=2,
                        content="Background details with [ref1] citation.",
                    )
                ],
                citations=["ref1"],
            ),
            DocumentSection(
                id="sec2",
                heading="Findings",
                level=1,
                content="Key findings are presented here.",
            ),
            DocumentSection(
                id="sec3",
                heading="Conclusion",
                level=1,
                content="Final conclusions.",
            ),
        ],
        bibliography=[
            {"key": "ref1", "title": "Source A", "author": "Author A", "year": "2024"},
            {"key": "ref2", "title": "Source B", "author": "Author B", "url": "https://example.com"},
        ],
    )


@pytest.fixture
def output_path(tmp_path):
    return tmp_path / "test_output.docx"


class TestExportToDocx:
    def test_creates_file(self, sample_content_tree, output_path):
        result = export_to_docx(sample_content_tree, output_path)
        assert result == output_path
        assert output_path.exists()
        assert output_path.stat().st_size > 0

    def test_creates_parent_directories(self, sample_content_tree, tmp_path):
        nested_path = tmp_path / "sub" / "dir" / "test.docx"
        export_to_docx(sample_content_tree, nested_path)
        assert nested_path.exists()

    def test_contains_title(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs]
        assert any("Test Report" in t for t in texts)

    def test_contains_abstract(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs]
        assert any("test abstract" in t for t in texts)

    def test_contains_headings(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        headings = [
            p.text for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        assert "Introduction" in headings
        assert "Background" in headings
        assert "Findings" in headings
        assert "Bibliography" in headings

    def test_contains_section_content(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs]
        all_text = " ".join(texts)
        assert "Introduction paragraph one" in all_text
        assert "Introduction paragraph two" in all_text

    def test_contains_bibliography_entries(self, sample_content_tree, output_path):
        export_to_docx(sample_content_tree, output_path)
        doc = Document(str(output_path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "[ref1]" in all_text
        assert "Author A" in all_text
        assert "Source A" in all_text

    def test_sanitizes_control_characters(self, tmp_path):
        tree = DocumentContentTree(
            doc_title="Control\x00Char Test",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Test\x0bHeading", level=1,
                    content="Content with\x08bad chars.",
                )
            ],
        )
        path = tmp_path / "sanitized.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "\x00" not in all_text
        assert "\x0b" not in all_text

    def test_empty_bibliography(self, tmp_path):
        tree = DocumentContentTree(
            doc_title="No Bib",
            doc_type="report",
            sections=[
                DocumentSection(id="s1", heading="Intro", level=1, content="Content."),
            ],
            bibliography=[],
        )
        path = tmp_path / "no_bib.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        headings = [
            p.text for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        assert "Bibliography" not in headings

    def test_no_abstract(self, tmp_path):
        tree = DocumentContentTree(
            doc_title="No Abstract",
            doc_type="report",
            abstract=None,
            sections=[
                DocumentSection(id="s1", heading="Intro", level=1, content="Content."),
            ],
        )
        path = tmp_path / "no_abstract.docx"
        export_to_docx(tree, path)
        assert path.exists()

    def test_deeply_nested_sections(self, tmp_path):
        tree = DocumentContentTree(
            doc_title="Nested",
            doc_type="report",
            sections=[
                DocumentSection(
                    id="s1", heading="Level 1", level=1,
                    content="L1 content.",
                    subsections=[
                        DocumentSection(
                            id="s1a", heading="Level 2", level=2,
                            content="L2 content.",
                            subsections=[
                                DocumentSection(
                                    id="s1a1", heading="Level 3", level=3,
                                    content="L3 content.",
                                )
                            ],
                        )
                    ],
                )
            ],
        )
        path = tmp_path / "nested.docx"
        export_to_docx(tree, path)
        doc = Document(str(path))
        headings = [
            p.text for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        assert "Level 1" in headings
        assert "Level 2" in headings
        assert "Level 3" in headings
