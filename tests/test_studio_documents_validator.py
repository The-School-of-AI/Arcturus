"""Tests for core/studio/documents/validator.py — DOCX and PDF validation."""

import pytest
from pathlib import Path

from core.schemas.studio_schema import DocumentContentTree, DocumentSection
from core.studio.documents.exporter_docx import export_to_docx
from core.studio.documents.exporter_pdf import export_to_pdf
from core.studio.documents.validator import validate_docx, validate_pdf

# Check if WeasyPrint native dependencies are available
_weasyprint_available = True
try:
    from weasyprint import HTML
except (ImportError, OSError):
    _weasyprint_available = False

_skip_no_weasyprint = pytest.mark.skipif(
    not _weasyprint_available,
    reason="WeasyPrint native libraries not available (gobject/pango/etc.)",
)


@pytest.fixture
def sample_content_tree():
    return DocumentContentTree(
        doc_title="Validation Test",
        doc_type="report",
        abstract="Test abstract for validation.",
        sections=[
            DocumentSection(
                id="sec1",
                heading="Introduction",
                level=1,
                content="Introduction content.",
                subsections=[
                    DocumentSection(
                        id="sec1a",
                        heading="Background",
                        level=2,
                        content="Background content.",
                    )
                ],
            ),
            DocumentSection(
                id="sec2",
                heading="Conclusion",
                level=1,
                content="Final conclusions.",
            ),
        ],
        bibliography=[
            {"key": "ref1", "title": "Source", "author": "Author"},
        ],
    )


class TestValidateDocx:
    def test_valid_docx(self, sample_content_tree, tmp_path):
        path = tmp_path / "valid.docx"
        export_to_docx(sample_content_tree, path)
        result = validate_docx(path, sample_content_tree)
        assert result["valid"] is True
        assert result["format"] == "docx"
        assert result["paragraph_count"] > 0
        assert result["heading_count"] > 0
        assert result["text_present"] is True

    def test_valid_docx_has_bibliography(self, sample_content_tree, tmp_path):
        path = tmp_path / "with_bib.docx"
        export_to_docx(sample_content_tree, path)
        result = validate_docx(path, sample_content_tree)
        assert result["bibliography_present"] is True

    def test_invalid_file_path(self, tmp_path):
        path = tmp_path / "nonexistent.docx"
        result = validate_docx(path)
        assert result["valid"] is False
        assert len(result["errors"]) > 0

    def test_corrupt_file(self, tmp_path):
        path = tmp_path / "corrupt.docx"
        path.write_text("not a docx file")
        result = validate_docx(path)
        assert result["valid"] is False

    def test_without_content_tree(self, sample_content_tree, tmp_path):
        path = tmp_path / "no_tree.docx"
        export_to_docx(sample_content_tree, path)
        result = validate_docx(path)
        assert result["valid"] is True

    def test_warns_on_missing_headings(self, tmp_path):
        """Validates warning when content_tree expects more headings."""
        tree = DocumentContentTree(
            doc_title="Many Sections",
            doc_type="report",
            sections=[
                DocumentSection(id=f"s{i}", heading=f"Section {i}", level=1, content=f"Content {i}.")
                for i in range(10)
            ],
        )
        # Export with only a few sections rendered
        from docx import Document
        doc = Document()
        doc.add_heading("Test", level=0)
        doc.add_heading("Only One", level=1)
        doc.add_paragraph("Some content.")
        path = tmp_path / "few_headings.docx"
        doc.save(str(path))

        result = validate_docx(path, tree)
        assert result["valid"] is True  # still valid, just has warnings
        assert len(result["warnings"]) > 0


@_skip_no_weasyprint
class TestValidatePdf:
    def test_valid_pdf(self, sample_content_tree, tmp_path):
        path = tmp_path / "valid.pdf"
        export_to_pdf(sample_content_tree, path)
        result = validate_pdf(path, sample_content_tree)
        assert result["valid"] is True
        assert result["format"] == "pdf"
        assert result["page_count"] >= 1
        assert result["text_present"] is True

    def test_valid_pdf_has_bibliography(self, sample_content_tree, tmp_path):
        path = tmp_path / "with_bib.pdf"
        export_to_pdf(sample_content_tree, path)
        result = validate_pdf(path, sample_content_tree)
        assert result["bibliography_present"] is True

    def test_invalid_file_path(self, tmp_path):
        path = tmp_path / "nonexistent.pdf"
        result = validate_pdf(path)
        assert result["valid"] is False
        assert len(result["errors"]) > 0

    def test_corrupt_file(self, tmp_path):
        path = tmp_path / "corrupt.pdf"
        path.write_text("not a pdf file")
        result = validate_pdf(path)
        assert result["valid"] is False

    def test_without_content_tree(self, sample_content_tree, tmp_path):
        path = tmp_path / "no_tree.pdf"
        export_to_pdf(sample_content_tree, path)
        result = validate_pdf(path)
        assert result["valid"] is True
